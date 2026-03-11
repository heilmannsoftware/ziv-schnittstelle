#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Word (.docx) documentation for ZIV Datenmodell 4.0"""

import os
import re
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# Reuse parsers from generate_docs.py
# ============================================================
import importlib.util
spec = importlib.util.spec_from_file_location("generate_docs", os.path.join(_BASE_DIR, "generate_docs.py"))
gen_docs = importlib.util.module_from_spec(spec)
spec.loader.exec_module(gen_docs)

parse_schema = gen_docs.parse_schema
parse_enums = gen_docs.parse_enums
parse_relations_from_dbml = gen_docs.parse_relations_from_dbml
parse_dbml_notes = gen_docs.parse_dbml_notes

# ============================================================
# Table groups (same as HTML docs)
# ============================================================
GROUPS = [
    ("Stammdaten/Verwaltung", ['Meta', 'Kehrbezirk', 'Mitarbeiter', 'Vertreter', 'Kommunikation', 'Fremdfirma', 'Untere_Aufsichtsbehoerde']),
    ("Standort", ['Ort', 'Strasse', 'Grundstueck', 'Grundstueck_Eigentuemer', 'Gebaeude_Gebaeudeteil', 'Gebaeude_Eigentuemer', 'Gebaeudeeigenschaft']),
    ("Abgasanlagen", ['Abgasanlage', 'Abschnitt', 'Schicht', 'Muendung', 'Reinigungsart']),
    ("Lueftungsanlagen", ['Lueftungsanlage', 'Raumlueftung', 'Raumlueftungsleitung', 'Raumluftoeffnung', 'Mess_Pruefergebnis_Raumluftoeffnung', 'Zubehoer']),
    ("Nutzung/Wohnung/Raum", ['Nutzungseinheit', 'Wohnung', 'Raum']),
    ("Feuerstaetten", ['Feuerstaette', 'Gasfeuerstaette', 'Oelfeuerstaette', 'Feste_Brennstoff_Feuerstaette', 'Sonderfeuerstaette', 'Waermepumpe', 'Feste_Brennstoff_Ableitbedingungen', 'Nachgeschalteter_Waermeaustauscher', 'Brennstoffversorgungsanlage', 'Raeucheranlage']),
    ("Messergebnisse", ['Mess_Pruefergebnis_Gasfeuerstaette', 'Mess_Pruefergebnis_44BImSchV_Gasfeuerstaette', 'Mess_Pruefergebnis_Oelfeuerstaette', 'Mess_Pruefergebnis_44BImSchV_Oelfeuerstaette', 'Mess_Pruefergebnis_Feststofffeuerstaette_HK', 'Mess_Pruefergebnis_Feststofffeuerstaette_ERF', 'Mess_Pruefergebnis_Sonderfeuerstaette', 'Mess_Pruefergebnis_Waermepumpe', 'Messgeraet', 'Messgeraet_44', 'Messunsicherheit']),
    ("Dunstabzug", ['Dunstabzuganlage_Leitung', 'Dunstabzugsanlage_Haube', 'Dunstabzugsanlage_Feuerstaette', 'Abschnitt_Dunstabzugsleitung', 'Ventilator_Dunstabzug', 'Mess_Pruefergebnis_Dunstabzugsleitung']),
    ("Verbrennungsluft", ['Verbrennungsluftzufuhr_Element', 'Feuerst_Verbr_Element']),
    ("GEG/Effizienz", ['GEG', 'Pruefergebnis_GEG', 'Effizienzlabel']),
    ("Hoheitliche Taetigkeiten", ['Abnahme', 'Anlassbezogene_Ueberpruefung', 'Sonstige_Arbeit', 'Feuerstaettenschau', 'Pruefergebnis_Feuerstaettenschau']),
    ("Feuerstaettenbescheid", ['Feuerstaettenbescheid_Grunddaten', 'Bescheid_Position', 'Bescheid_Termin_Position']),
    ("Abrechnung", ['Leistung', 'Aufteilung_Leistung', 'Datum_Ausfuehrung_Leistung', 'Rechnung', 'Landesrechtliche_Vorschrift', 'Bundes_KUEO']),
    ("Kehrbuch", ['Kehrbuch', 'Kehrbuch_Gebaeude', 'Kehrbuch_Gebaeude_Eigentuemer', 'Kehrbuch_Nutzungseinheit', 'Kehrbuch_Anlage', 'Kehrbuch_Taetigkeit', 'Kehrbuch_Abnahme', 'Kehrbuch_Anlassbezogene_Ueberpruefung']),
    ("Berichte/Maengel", ['Bericht', 'Mangel']),
    ("Sonstiges", ['Dachskizze', 'Dachskizze_Element', 'Belegungsplan', 'Dokument', 'Gefaehrdungsbeurteilung']),
]

# ============================================================
# Colors
# ============================================================
COLOR_ACCENT = RGBColor(0, 113, 227)
COLOR_PK = RGBColor(180, 140, 0)
COLOR_FK = RGBColor(0, 113, 227)
COLOR_HEADER_BG = 'D6E4F0'
COLOR_PK_ROW_BG = 'FFF8E1'
COLOR_ALT_ROW_BG = 'F5F5F7'
COLOR_GROUP_HEADER = RGBColor(29, 29, 31)
COLOR_GRAY = RGBColor(110, 110, 115)

# ============================================================
# Helpers
# ============================================================
def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph for internal linking."""
    bookmark_name = re.sub(r'[^A-Za-z0-9_]', '_', bookmark_name)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="0" w:name="{bookmark_name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="0"/>')
    tag.addprevious(start)
    tag.addnext(end)


def add_hyperlink(paragraph, text, bookmark_name, color=COLOR_ACCENT):
    """Add an internal hyperlink to a bookmark."""
    bookmark_name = re.sub(r'[^A-Za-z0-9_]', '_', bookmark_name)

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} w:anchor="{bookmark_name}"/>'
    )
    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:color w:val="{str(color)}"/>'
        f'<w:u w:val="single"/><w:sz w:val="18"/><w:szCs w:val="18"/></w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def format_table_style(table):
    """Apply consistent formatting to a Word table."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


# ============================================================
# Build Word document
# ============================================================
def generate_word(tables, enums, relations, notes, arts):
    doc = Document()

    # ---- Page setup: Landscape A4 ----
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ---- Default font ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9)

    # ---- Heading styles ----
    for level in [1, 2, 3]:
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = COLOR_GROUP_HEADER

    doc.styles['Heading 1'].font.size = Pt(24)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(12)

    # Build relation maps
    parent_refs = {}
    child_refs = {}
    for rel in relations:
        child = rel['kind_tabelle']
        parent = rel['eltern_tabelle']
        fk = rel['fk_feld']
        kard = rel['kardinalitaet']
        comment = rel['kommentar']
        parent_refs.setdefault(child, []).append((fk, parent, comment, kard))
        child_refs.setdefault(parent, []).append((child, fk, comment, kard))

    # Enum value set mapping
    enum_value_sets = {}
    for enum_name, values in enums.items():
        key = frozenset(v['wert'] for v in values)
        if key:
            enum_value_sets[key] = enum_name

    def find_enum_for_check(check_values):
        if not check_values:
            return None
        return enum_value_sets.get(frozenset(check_values))

    # Enum usage
    enum_usage = {}
    for tname, tdata in tables.items():
        for col in tdata['columns']:
            if col['check_values']:
                enum_name = find_enum_for_check(col['check_values'])
                if enum_name:
                    enum_usage.setdefault(enum_name, []).append((tname, col['name']))

    # ---- Stats ----
    total_tables = len(tables)
    total_relations = len(relations)
    total_enums = len(enums)
    total_enum_values = sum(len(v) for v in enums.values())

    # ============================================================
    # Title page
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ZIV Datenmodell 4.0')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_GROUP_HEADER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Dokumentation der SQLite-Datenbank')
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats.add_run(f'{total_tables} Tabellen  |  {total_relations} Relationen  |  {total_enums} Enums  |  {total_enum_values:,} Enum-Werte')
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    gen_date = doc.add_paragraph()
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_date.add_run(f'Generiert am {date.today().isoformat()}')
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_GRAY

    doc.add_page_break()

    # ============================================================
    # Table of Contents
    # ============================================================
    toc_heading = doc.add_heading('Inhaltsverzeichnis', level=1)
    add_bookmark(toc_heading, 'TOC')

    for gname, gtables in GROUPS:
        # Group heading in TOC
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        p.space_after = Pt(2)
        run = p.add_run(gname)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GROUP_HEADER

        for tname in gtables:
            if tname not in tables:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.space_before = Pt(1)
            p.space_after = Pt(1)
            add_hyperlink(p, tname, f'table_{tname}')

    # Enum TOC entry
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run('Enums (Stammdaten-Wertetabellen)')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GROUP_HEADER

    sorted_enums = sorted(enums.keys())
    for enum_name in sorted_enums:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        add_hyperlink(p, enum_name, f'enum_{enum_name}')

    # ============================================================
    # Tables by group
    # ============================================================
    for gname, gtables in GROUPS:
        doc.add_page_break()

        group_heading = doc.add_heading(gname, level=1)
        add_bookmark(group_heading, f'group_{gname.replace("/", "_").replace(" ", "_")}')

        for tname in gtables:
            if tname not in tables:
                continue
            tdata = tables[tname]
            tnotes = notes.get(tname, {})
            tarts = arts.get(tname, {})

            doc.add_page_break()

            # Table heading
            table_heading = doc.add_heading(tname, level=2)
            add_bookmark(table_heading, f'table_{tname}')

            # Column count info
            p = doc.add_paragraph()
            run = p.add_run(f'{len(tdata["columns"])} Spalten')
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GRAY

            # ---- Columns table ----
            col_headers = ['Spalte', 'Typ', 'PK', 'FK', 'NOT NULL', 'Default', 'Art', 'Beschreibung']
            num_cols = len(col_headers)

            tbl = doc.add_table(rows=1, cols=num_cols)
            format_table_style(tbl)

            # Set column widths
            widths = [Cm(5.5), Cm(1.8), Cm(1.0), Cm(1.0), Cm(1.5), Cm(2.0), Cm(1.8), Cm(12.0)]
            for i, width in enumerate(widths):
                tbl.columns[i].width = width

            # Header row
            hdr = tbl.rows[0]
            for i, h_text in enumerate(col_headers):
                cell = hdr.cells[i]
                cell.text = h_text
                set_cell_shading(cell, COLOR_HEADER_BG)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(50, 50, 50)

            # Data rows
            for idx, col in enumerate(tdata['columns']):
                row = tbl.add_row()

                # Alternating row colors + PK highlight
                if col['pk']:
                    bg = COLOR_PK_ROW_BG
                elif idx % 2 == 1:
                    bg = COLOR_ALT_ROW_BG
                else:
                    bg = None

                if bg:
                    for cell in row.cells:
                        set_cell_shading(cell, bg)

                # Spalte
                cell = row.cells[0]
                p = cell.paragraphs[0]
                run = p.add_run(col['name'])
                run.font.size = Pt(8)
                run.font.bold = col['pk']

                # Typ
                row.cells[1].paragraphs[0].add_run(col['type']).font.size = Pt(8)

                # PK
                if col['pk']:
                    run = row.cells[2].paragraphs[0].add_run('PK')
                    run.font.size = Pt(8)
                    run.font.bold = True
                    run.font.color.rgb = COLOR_PK

                # FK
                if col['fk']:
                    run = row.cells[3].paragraphs[0].add_run('FK')
                    run.font.size = Pt(8)
                    run.font.bold = True
                    run.font.color.rgb = COLOR_FK

                # NOT NULL
                if col['not_null'] or col['pk']:
                    row.cells[4].paragraphs[0].add_run('Ja').font.size = Pt(8)

                # Default
                if col['default']:
                    row.cells[5].paragraphs[0].add_run(str(col['default'])).font.size = Pt(8)

                # Art
                art_val = tarts.get(col['name'], '')
                if art_val:
                    run = row.cells[6].paragraphs[0].add_run(art_val)
                    run.font.size = Pt(8)
                    if art_val.startswith('§ 19'):
                        run.font.color.rgb = RGBColor(0, 140, 60)
                        run.font.bold = True
                    elif art_val == 'QS':
                        run.font.color.rgb = COLOR_ACCENT
                        run.font.bold = True
                    elif art_val == 'NEIN':
                        run.font.color.rgb = RGBColor(200, 0, 0)

                # Beschreibung
                desc_cell = row.cells[7]
                desc_p = desc_cell.paragraphs[0]

                note = tnotes.get(col['name'], '')
                if note:
                    run = desc_p.add_run(note)
                    run.font.size = Pt(8)
                    run.font.color.rgb = COLOR_GRAY

                # FK reference
                if col['fk']:
                    ref_table = col['fk']['ref_table']
                    if note:
                        desc_p.add_run('\n').font.size = Pt(8)
                    run = desc_p.add_run(f'\u2192 ')
                    run.font.size = Pt(8)
                    add_hyperlink(desc_p, ref_table, f'table_{ref_table}')

                # Enum reference
                if col['check_values']:
                    enum_name = find_enum_for_check(col['check_values'])
                    if enum_name:
                        if note or col['fk']:
                            desc_p.add_run('\n').font.size = Pt(8)
                        run = desc_p.add_run('Enum: ')
                        run.font.size = Pt(8)
                        add_hyperlink(desc_p, enum_name, f'enum_{enum_name}', RGBColor(36, 138, 61))

            # ---- Relations summary ----
            parents = parent_refs.get(tname, [])
            children = child_refs.get(tname, [])

            if parents or children:
                doc.add_paragraph()  # spacer

                if parents:
                    p = doc.add_paragraph()
                    run = p.add_run('Referenziert: ')
                    run.font.bold = True
                    run.font.size = Pt(9)
                    for i, (fk, parent, comment, kard) in enumerate(parents):
                        if i > 0:
                            p.add_run(', ').font.size = Pt(9)
                        add_hyperlink(p, parent, f'table_{parent}')
                        run = p.add_run(f' ({fk}, {kard})')
                        run.font.size = Pt(8)
                        run.font.color.rgb = COLOR_GRAY

                if children:
                    p = doc.add_paragraph()
                    run = p.add_run('Referenziert von: ')
                    run.font.bold = True
                    run.font.size = Pt(9)
                    for i, (child, fk, comment, kard) in enumerate(children):
                        if i > 0:
                            p.add_run(', ').font.size = Pt(9)
                        add_hyperlink(p, child, f'table_{child}')
                        run = p.add_run(f' ({fk}, {kard})')
                        run.font.size = Pt(8)
                        run.font.color.rgb = COLOR_GRAY

    # ============================================================
    # Enums section
    # ============================================================
    doc.add_page_break()
    enum_heading = doc.add_heading('Enums (Stammdaten-Wertetabellen)', level=1)
    add_bookmark(enum_heading, 'enums_section')

    p = doc.add_paragraph()
    run = p.add_run(f'{total_enums} Enums mit {total_enum_values:,} Werten')
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_GRAY

    for enum_name in sorted_enums:
        values = enums[enum_name]
        usage = enum_usage.get(enum_name, [])

        doc.add_page_break()

        eh = doc.add_heading(f'{enum_name} ({len(values)} Werte)', level=2)
        add_bookmark(eh, f'enum_{enum_name}')

        # Usage info
        if usage:
            p = doc.add_paragraph()
            run = p.add_run('Verwendet in: ')
            run.font.bold = True
            run.font.size = Pt(9)
            for i, (ut, uc) in enumerate(usage):
                if i > 0:
                    p.add_run(', ').font.size = Pt(9)
                add_hyperlink(p, f'{ut}.{uc}', f'table_{ut}')

        # Values table
        tbl = doc.add_table(rows=1, cols=2)
        format_table_style(tbl)
        tbl.columns[0].width = Cm(6)
        tbl.columns[1].width = Cm(20)

        # Header
        for i, h_text in enumerate(['Wert', 'Beschreibung']):
            cell = tbl.rows[0].cells[i]
            cell.text = h_text
            set_cell_shading(cell, COLOR_HEADER_BG)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Values
        for idx, v in enumerate(values):
            row = tbl.add_row()
            if idx % 2 == 1:
                for cell in row.cells:
                    set_cell_shading(cell, COLOR_ALT_ROW_BG)

            row.cells[0].paragraphs[0].add_run(v['wert']).font.size = Pt(8)
            row.cells[1].paragraphs[0].add_run(v['beschreibung']).font.size = Pt(8)

    return doc


# ============================================================
# Main
# ============================================================
if __name__ == '__main__':
    print("Parsing schema.sql...")
    tables = parse_schema(os.path.join(_BASE_DIR, 'schema.sql'))
    print(f"  Found {len(tables)} tables")

    print("Parsing stammdaten_enums.sql...")
    enums = parse_enums(os.path.join(_BASE_DIR, 'stammdaten_enums.sql'))
    print(f"  Found {len(enums)} enums")

    print("Parsing DBML for relations...")
    relations = parse_relations_from_dbml(os.path.join(_BASE_DIR, 'Datenmodell4_0.dbml'))
    print(f"  Found {len(relations)} relations")

    print("Parsing Datenmodell4_0.dbml for notes...")
    dbml_notes, dbml_arts = parse_dbml_notes(os.path.join(_BASE_DIR, 'Datenmodell4_0.dbml'))
    print(f"  Found notes for {len(dbml_notes)} tables")

    print("Generating Word document...")
    doc = generate_word(tables, enums, relations, dbml_notes, dbml_arts)

    output_path = os.path.join(_BASE_DIR, 'dokumentation.docx')
    doc.save(output_path)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"Done! Written to {output_path} ({size_kb:.0f} KB)")
